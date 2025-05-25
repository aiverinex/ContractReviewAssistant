#!/usr/bin/env python3
"""
Contract Reviewer Crew - Main Application Entry Point

A production-grade CrewAI-based contract review system that automates legal
contract analysis using specialized AI agents for clause detection, risk
analysis, and redline suggestions.

This application is designed for submission to the CrewAI Marketplace:
https://marketplace.crewai.com
"""

import os
import sys
import json
import traceback
from datetime import datetime
from pathlib import Path

# Document processing imports
import docx
import pdfplumber
from PyPDF2 import PdfReader

# Environment and configuration
from dotenv import load_dotenv
import yaml

# CrewAI components
from crew.crew import ContractReviewCrew

def load_environment():
    """Load environment variables and validate configuration."""
    load_dotenv()
    
    # Check for required API key
    openai_key = os.getenv("OPENAI_API_KEY")
    if not openai_key:
        print("❌ Error: OPENAI_API_KEY environment variable is required")
        print("Please set your OpenAI API key in the .env file")
        sys.exit(1)
    
    print("✅ Environment loaded successfully")
    return True

def load_config():
    """Load configuration from YAML file."""
    try:
        config_path = Path("config/config.yaml")
        if config_path.exists():
            with open(config_path, 'r') as file:
                config = yaml.safe_load(file)
                print("✅ Configuration loaded from config.yaml")
                return config
        else:
            print("⚠️  Warning: config.yaml not found, using defaults")
            return {}
    except Exception as e:
        print(f"⚠️  Warning: Failed to load config.yaml: {e}")
        return {}

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using pdfplumber with fallback to PyPDF2.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text content
    """
    text = ""
    
    try:
        # Primary method: pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if text.strip():
            print(f"✅ Text extracted from PDF using pdfplumber: {len(text)} characters")
            return text
            
    except Exception as e:
        print(f"⚠️  pdfplumber extraction failed: {e}")
    
    try:
        # Fallback method: PyPDF2
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        if text.strip():
            print(f"✅ Text extracted from PDF using PyPDF2: {len(text)} characters")
            return text
            
    except Exception as e:
        print(f"❌ PDF text extraction failed with both methods: {e}")
    
    return text

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text content
    """
    try:
        doc = docx.Document(file_path)
        text = ""
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        print(f"✅ Text extracted from DOCX: {len(text)} characters")
        return text
        
    except Exception as e:
        print(f"❌ DOCX text extraction failed: {e}")
        return ""

def load_contract_text(file_path: str = None) -> str:
    """
    Load contract text from file or use sample data.
    
    Args:
        file_path (str, optional): Path to contract file
        
    Returns:
        str: Contract text content
    """
    if file_path and Path(file_path).exists():
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
                    print(f"✅ Loaded text file: {len(text)} characters")
                    return text
            except Exception as e:
                print(f"❌ Failed to load text file: {e}")
                
        elif file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
            
        elif file_extension in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
            
        else:
            print(f"❌ Unsupported file format: {file_extension}")
    
    # Use sample contract if no file provided or file loading failed
    sample_path = Path("sample_data/sample_contract.txt")
    if sample_path.exists():
        try:
            with open(sample_path, 'r', encoding='utf-8') as file:
                text = file.read()
                print(f"✅ Using sample contract: {len(text)} characters")
                return text
        except Exception as e:
            print(f"❌ Failed to load sample contract: {e}")
    
    print("❌ No contract text available for processing")
    return ""

def save_review_results(results: dict, output_dir: str = "output") -> bool:
    """
    Save contract review results to files.
    
    Args:
        results (dict): Review results to save
        output_dir (str): Output directory path
        
    Returns:
        bool: Success status
    """
    try:
        # Create output directory
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Save JSON results
        json_file = output_path / f"review_summary_{timestamp}.json"
        with open(json_file, 'w', encoding='utf-8') as file:
            json.dump(results, file, indent=2, ensure_ascii=False)
        
        # Save human-readable summary
        txt_file = output_path / f"review_summary_{timestamp}.txt"
        with open(txt_file, 'w', encoding='utf-8') as file:
            file.write("CONTRACT REVIEW SUMMARY\n")
            file.write("=" * 50 + "\n\n")
            
            if results.get("success"):
                exec_summary = results.get("executive_summary", {})
                overview = exec_summary.get("contract_overview", {})
                assessment = exec_summary.get("overall_assessment", {})
                
                file.write(f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                
                file.write("KEY METRICS:\n")
                file.write(f"- Clauses Analyzed: {overview.get('clauses_analyzed', 0)}\n")
                file.write(f"- Risks Identified: {overview.get('risks_identified', 0)}\n")
                file.write(f"- High-Severity Risks: {overview.get('high_severity_risks', 0)}\n")
                file.write(f"- Redline Suggestions: {overview.get('redline_suggestions', 0)}\n")
                file.write(f"- Critical Changes Needed: {overview.get('critical_changes_needed', 0)}\n\n")
                
                file.write("OVERALL ASSESSMENT:\n")
                file.write(f"- Risk Level: {assessment.get('risk_level', 'Unknown')}\n")
                file.write(f"- Contract Quality: {assessment.get('contract_quality', 'Unknown')}\n")
                file.write(f"- Recommended Action: {assessment.get('recommended_action', 'Review needed')}\n\n")
                
                key_concerns = assessment.get('key_concerns', [])
                if key_concerns:
                    file.write("KEY CONCERNS:\n")
                    for concern in key_concerns:
                        file.write(f"- {concern}\n")
                    file.write("\n")
                
                next_steps = results.get("next_steps", [])
                if next_steps:
                    file.write("RECOMMENDED NEXT STEPS:\n")
                    for i, step in enumerate(next_steps, 1):
                        file.write(f"{i}. {step}\n")
            else:
                file.write(f"ERROR: {results.get('error', 'Unknown error occurred')}\n")
        
        print(f"✅ Results saved to:")
        print(f"   - {json_file}")
        print(f"   - {txt_file}")
        return True
        
    except Exception as e:
        print(f"❌ Failed to save results: {e}")
        return False

def print_executive_summary(results: dict):
    """Print a formatted executive summary to console."""
    if not results.get("success"):
        print(f"\n❌ CONTRACT REVIEW FAILED")
        print(f"Error: {results.get('error', 'Unknown error')}")
        return
    
    print("\n" + "=" * 60)
    print("🏛️  CONTRACT REVIEW EXECUTIVE SUMMARY")
    print("=" * 60)
    
    exec_summary = results.get("executive_summary", {})
    overview = exec_summary.get("contract_overview", {})
    assessment = exec_summary.get("overall_assessment", {})
    
    print(f"\n📊 KEY METRICS:")
    print(f"   • Clauses Analyzed: {overview.get('clauses_analyzed', 0)}")
    print(f"   • High-Importance Clauses: {overview.get('high_importance_clauses', 0)}")
    print(f"   • Risks Identified: {overview.get('risks_identified', 0)}")
    print(f"   • High-Severity Risks: {overview.get('high_severity_risks', 0)}")
    print(f"   • Redline Suggestions: {overview.get('redline_suggestions', 0)}")
    print(f"   • Critical Changes: {overview.get('critical_changes_needed', 0)}")
    
    print(f"\n🎯 OVERALL ASSESSMENT:")
    risk_level = assessment.get('risk_level', 'Unknown')
    risk_emoji = "🔴" if risk_level == "HIGH" else "🟡" if risk_level == "MEDIUM" else "🟢"
    print(f"   • Risk Level: {risk_emoji} {risk_level}")
    
    quality = assessment.get('contract_quality', 'Unknown')
    quality_emoji = "✅" if quality == "GOOD" else "⚠️" if quality == "FAIR" else "❌"
    print(f"   • Contract Quality: {quality_emoji} {quality}")
    print(f"   • Recommended Action: {assessment.get('recommended_action', 'Review needed')}")
    
    key_concerns = assessment.get('key_concerns', [])
    if key_concerns:
        print(f"\n⚠️  KEY CONCERNS:")
        for concern in key_concerns[:5]:  # Limit to top 5
            print(f"   • {concern}")
    
    next_steps = results.get("next_steps", [])
    if next_steps:
        print(f"\n📋 RECOMMENDED NEXT STEPS:")
        for i, step in enumerate(next_steps[:5], 1):  # Limit to top 5
            print(f"   {i}. {step}")

def main():
    """Main application entry point."""
    print("🏛️  Contract Reviewer Crew - AI-Powered Contract Analysis")
    print("=" * 60)
    print("A CrewAI-based system for automated legal contract review")
    print("Marketplace: https://marketplace.crewai.com")
    print("=" * 60)
    
    try:
        # Load environment and configuration
        load_environment()
        config = load_config()
        
        # Get contract file path from command line argument or use sample
        contract_file = None
        if len(sys.argv) > 1:
            contract_file = sys.argv[1]
            print(f"\n📄 Processing contract file: {contract_file}")
        else:
            print("\n📄 No file specified, using sample contract")
        
        # Load contract text
        print("\n🔄 Loading contract text...")
        contract_text = load_contract_text(contract_file)
        
        if not contract_text.strip():
            print("❌ No contract text available for processing")
            sys.exit(1)
        
        print(f"✅ Contract loaded: {len(contract_text)} characters")
        
        # Initialize CrewAI contract review system
        print("\n🤖 Initializing AI agents...")
        crew = ContractReviewCrew()
        
        # Check crew status
        status = crew.get_crew_status()
        if not status.get("crew_initialized"):
            print("❌ Failed to initialize CrewAI crew")
            sys.exit(1)
        
        if not status.get("openai_configured"):
            print("❌ OpenAI API key not configured")
            sys.exit(1)
        
        print("✅ AI agents initialized successfully")
        
        # Execute contract review
        print("\n🚀 Starting contract review process...")
        results = crew.review_contract(contract_text)
        
        if not results.get("success"):
            print(f"❌ Contract review failed: {results.get('error', 'Unknown error')}")
            sys.exit(1)
        
        # Display results
        print_executive_summary(results)
        
        # Save results to files
        print(f"\n💾 Saving results...")
        save_success = save_review_results(results)
        
        if save_success:
            print("\n✅ Contract review completed successfully!")
            print("📁 Check the 'output' directory for detailed results")
        else:
            print("\n⚠️  Contract review completed but failed to save results")
            
    except KeyboardInterrupt:
        print("\n\n⏹️  Contract review interrupted by user")
        sys.exit(1)
        
    except Exception as e:
        print(f"\n❌ Unexpected error: {str(e)}")
        print(f"\nFull traceback:\n{traceback.format_exc()}")
        sys.exit(1)

if __name__ == "__main__":
    main()
